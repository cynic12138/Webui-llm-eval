#!/usr/bin/env python3
"""生成 Webui-LLM-Eval 项目使用说明文档 (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime


def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color
    })
    shading.append(shading_elm)


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加样式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2F5496')

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()  # 空行
    return table


def main():
    doc = Document()

    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ========== 样式设置 ==========
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    for i in range(1, 5):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

    # ========== 封面 ==========
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Webui-LLM-Eval')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('大语言模型评测平台 — 详细使用说明文档')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'文档版本：V1.0\n生成日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # ========== 目录页 ==========
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 项目概述',
        '2. 系统架构',
        '3. 环境要求与依赖',
        '4. 安装与部署',
        '    4.1 Docker Compose 部署（推荐）',
        '    4.2 一键脚本部署',
        '    4.3 手动部署',
        '    4.4 开发模式启动',
        '5. 配置说明',
        '6. 功能模块详解',
        '    6.1 用户认证与权限',
        '    6.2 模型管理',
        '    6.3 数据集管理',
        '    6.4 基准测试库',
        '    6.5 评测任务',
        '    6.6 结果分析与可视化',
        '    6.7 报告导出',
        '    6.8 ELO 排行榜',
        '    6.9 模型竞技场',
        '    6.10 Prompt 管理',
        '    6.11 AI 助手',
        '    6.12 管理后台',
        '7. 评测维度详解',
        '8. API 接口概览',
        '9. 常见问题与故障排查',
        '10. 扩展开发指南',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # ========== 1. 项目概述 ==========
    doc.add_heading('1. 项目概述', level=1)

    doc.add_paragraph(
        'Webui-LLM-Eval 是一个企业级大语言模型（LLM）综合评测平台，旨在为研究人员和工程团队提供'
        '标准化、可复现的模型评估能力。平台通过 Web 界面提供全流程的模型评测服务，支持多模型并行评测、'
        '20+ 标准基准测试、多维度评估指标、可视化报告生成以及 ELO 排名系统。'
    )

    doc.add_heading('核心特性', level=2)
    features = [
        ('多模型并行评测', '同时对多个 LLM 进行相同数据集的评测，支持横向对比'),
        ('20+ 标准基准测试', '内置 MMLU-Pro、GSM8K、HumanEval、C-Eval、TruthfulQA、HealthBench 等'),
        ('LLM-as-Judge 智能评分', '利用大语言模型作为评判者进行智能打分，内置偏差消除机制'),
        ('客观指标评估', '支持 ROUGE、BLEU、METEOR、BERTScore、Exact Match、Token F1、余弦相似度'),
        ('幻觉检测', '通过一致性采样检测模型生成内容的幻觉问题'),
        ('鲁棒性测试', '对输入进行扰动测试，评估模型的稳定性'),
        ('RAG 评估', '检索增强生成的专项评测，包含相关性和忠实度分析'),
        ('代码执行验证', '在 Docker 沙箱中执行代码并验证测试用例'),
        ('链式推理评估', 'Chain-of-Thought 推理质量评估'),
        ('多轮对话测试', '评估模型在多轮对话场景下的表现'),
        ('多语言能力测试', '跨语言能力评估'),
        ('结构化输出验证', 'JSON/XML 等结构化输出的 Schema 验证'),
        ('函数调用测试', '工具调用/Function Calling 能力评估'),
        ('实时进度追踪', '基于 WebSocket 的实时进度更新'),
        ('可视化报告', '雷达图、柱状图、分布图等多维可视化'),
        ('报告导出', '支持 PDF、Excel、JSONL 格式导出'),
        ('ELO 排名系统', '基于对战记录的动态排行榜'),
        ('团队协作', '多用户支持，基于角色的权限控制'),
        ('Prompt 管理', 'Prompt 版本控制与实验追踪'),
        ('AI 助手', '内置智能助手辅助解读评测结果'),
    ]
    for name, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ========== 2. 系统架构 ==========
    doc.add_heading('2. 系统架构', level=1)

    doc.add_paragraph(
        '本系统采用前后端分离的微服务架构，共包含 6 层核心组件：'
    )

    arch_items = [
        ('Nginx 反向代理', '统一入口，将 /api 路由至后端，/ 路由至前端'),
        ('前端 (Frontend)', 'Next.js 14 + TypeScript + React 18 + Ant Design 5，提供完整的 Web 界面'),
        ('后端 (Backend)', 'FastAPI + Python 3.11，提供 RESTful API，包含 18 个路由模块'),
        ('任务队列 (Celery)', 'Celery + Redis 异步任务处理，负责评测任务的调度与执行'),
        ('评测引擎 (Eval Engine)', '核心评测逻辑，包含 20+ 评测器和多个 LLM Provider 适配器'),
        ('基础设施', 'PostgreSQL（数据存储）+ Redis（缓存/消息队列）+ MinIO（对象存储）'),
    ]
    for name, desc in arch_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{name}：')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('请求流程', level=2)
    doc.add_paragraph(
        '用户请求 → Nginx → 前端页面 / 后端 API → Celery 异步任务 → 评测引擎 → '
        'LLM Provider API → 评测结果存储 → WebSocket 实时推送 → 前端展示'
    )

    doc.add_heading('项目目录结构', level=2)
    structure = [
        ('frontend/', '前端应用（Next.js 14 + TypeScript）'),
        ('  src/app/(auth)/', '认证页面：登录、注册、登出'),
        ('  src/app/(dashboard)/', '主要功能页面：仪表盘、评测、模型、数据集等'),
        ('  src/components/', 'React 组件库（图表、评测、布局等）'),
        ('  src/lib/', '工具函数与 Hooks（API 客户端、WebSocket、认证）'),
        ('  src/types/', 'TypeScript 类型定义'),
        ('backend/', '后端应用（FastAPI + Python）'),
        ('  main.py', 'FastAPI 应用入口'),
        ('  app/api/v1/', 'RESTful API 路由（18 个模块）'),
        ('  app/core/', '核心配置（Celery、安全、依赖注入）'),
        ('  app/db/', '数据库模型与会话管理（20+ 表）'),
        ('  app/schemas/', 'Pydantic 请求/响应模型'),
        ('  app/services/', '业务服务层（评测、报告、存储、通知等）'),
        ('  alembic/', '数据库迁移脚本'),
        ('eval_engine/', '评测引擎核心逻辑'),
        ('  evaluators/', '20+ 评测模块'),
        ('  providers/', 'LLM API 适配器（OpenAI、Anthropic 等）'),
        ('  benchmark_data/', '内置基准测试数据集（JSONL）'),
        ('  engine.py', '评测引擎主调度器'),
        ('nginx/', 'Nginx 反向代理配置'),
        ('docker-compose.yml', 'Docker 多容器编排'),
        ('.env.example', '环境变量配置模板'),
        ('setup.sh / setup.ps1', '自动部署脚本（Linux/Windows）'),
        ('ecosystem.config.js', 'PM2 进程管理配置'),
    ]
    add_styled_table(doc, ['路径', '说明'], structure, [2.5, 4.0])

    doc.add_page_break()

    # ========== 3. 环境要求与依赖 ==========
    doc.add_heading('3. 环境要求与依赖', level=1)

    doc.add_heading('系统环境要求', level=2)
    env_reqs = [
        ('Python', '3.10 - 3.11', '推荐 3.11，与 PyTorch 兼容性最佳'),
        ('Node.js', '18+', '推荐 20+'),
        ('PostgreSQL', '14+', '推荐 15+，用于数据持久化'),
        ('Redis', '6+', '推荐 7+，用于缓存和消息队列'),
        ('MinIO', '最新版', '对象存储，用于文件/报告存储'),
        ('Docker', '20+（可选）', 'Docker Compose 部署方式需要'),
        ('PM2', '最新版（可选）', '进程管理，手动部署方式需要'),
    ]
    add_styled_table(doc, ['组件', '版本要求', '说明'], env_reqs, [1.5, 1.5, 3.5])

    doc.add_heading('后端主要依赖（Python）', level=2)
    py_deps = [
        ('FastAPI 0.110.0', 'Web 框架'),
        ('SQLAlchemy 2.0.28', 'ORM 数据库操作'),
        ('Celery 5.3.6', '异步任务队列'),
        ('Redis 5.0.2', '缓存客户端'),
        ('MinIO 7.2.5', '对象存储客户端'),
        ('OpenAI 1.14.0', 'OpenAI API SDK'),
        ('Anthropic 0.20.0', 'Anthropic API SDK'),
        ('Transformers 4.38.2', 'Hugging Face 模型库'),
        ('PyTorch 2.2.1', '深度学习框架'),
        ('ROUGE-Score / SacreBLEU', '文本评估指标'),
        ('Detoxify 0.5.2', '毒性检测'),
        ('ReportLab / OpenPyXL', '报告生成（PDF/Excel）'),
    ]
    add_styled_table(doc, ['依赖', '用途'], py_deps, [2.5, 4.0])

    doc.add_heading('前端主要依赖（Node.js）', level=2)
    js_deps = [
        ('Next.js 14.2.3', 'React 全栈框架'),
        ('React 18', 'UI 组件库基础'),
        ('TypeScript', '类型安全'),
        ('Ant Design 5.15.0', 'UI 组件库'),
        ('ECharts 5.5.0', '数据可视化图表'),
        ('Axios 1.6.8', 'HTTP 请求库'),
        ('react-markdown 9.0.1', 'Markdown 渲染'),
    ]
    add_styled_table(doc, ['依赖', '用途'], js_deps, [2.5, 4.0])

    doc.add_page_break()

    # ========== 4. 安装与部署 ==========
    doc.add_heading('4. 安装与部署', level=1)

    doc.add_heading('4.1 Docker Compose 部署（推荐）', level=2)
    doc.add_paragraph('Docker Compose 是最简单的部署方式，自动编排所有服务。')

    steps = [
        '克隆项目仓库并进入目录',
        '复制配置文件：cp .env.example .env',
        '编辑 .env 文件，修改数据库密码、密钥等敏感配置',
        '启动所有服务：docker-compose up -d',
        '访问前端界面：http://localhost:3000',
        '访问 API 文档：http://localhost:8000/docs',
        '访问 MinIO 控制台：http://localhost:9001',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph('Docker Compose 将启动以下 6 个服务：')
    services = [
        ('postgres', 'PostgreSQL 15', '5432', '数据库'),
        ('redis', 'Redis 7', '6379', '缓存/消息队列'),
        ('minio', 'MinIO', '9000/9001', '对象存储'),
        ('backend', 'FastAPI', '8000', '后端 API'),
        ('celery', 'Celery Worker', '-', '异步任务处理'),
        ('frontend', 'Next.js', '3000', '前端界面'),
    ]
    add_styled_table(doc, ['服务名', '镜像', '端口', '说明'], services, [1.2, 1.5, 1.3, 2.5])

    doc.add_heading('4.2 一键脚本部署', level=2)
    doc.add_paragraph(
        '项目提供了一键部署脚本，适用于 Linux/macOS（setup.sh）和 Windows（setup.ps1）系统。'
    )
    steps = [
        '赋予脚本执行权限：chmod +x setup.sh',
        '首次运行脚本：./setup.sh（将自动生成 .env 配置文件）',
        '编辑 .env 文件，配置数据库连接、API 密钥等参数',
        '再次运行脚本：./setup.sh（将安装依赖并启动所有服务）',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph(
        '脚本将自动完成：环境检查、Python 虚拟环境创建、依赖安装、数据库初始化、'
        '前端构建、PM2 进程启动等全部操作。'
    )

    doc.add_heading('4.3 手动部署', level=2)

    doc.add_heading('Step 1：创建 PostgreSQL 数据库', level=3)
    doc.add_paragraph(
        'sudo -u postgres psql -c "CREATE USER llmeval WITH PASSWORD \'llmeval123\';"'
    )
    doc.add_paragraph(
        'sudo -u postgres psql -c "CREATE DATABASE llmeval OWNER llmeval;"'
    )

    doc.add_heading('Step 2：配置环境变量', level=3)
    doc.add_paragraph('cp .env.example .env    # 复制配置模板并编辑')

    doc.add_heading('Step 3：安装后端', level=3)
    manual_steps = [
        'cd backend',
        'python3.11 -m venv .venv',
        'source .venv/bin/activate',
        'pip install -r requirements.txt',
    ]
    for step in manual_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Step 4：安装前端', level=3)
    manual_steps = [
        'cd frontend',
        'npm ci',
        'cp .env.local.example .env.local    # 配置 API 地址',
        'npm run build',
    ]
    for step in manual_steps:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Step 5：启动服务', level=3)
    doc.add_paragraph('pm2 start ecosystem.config.js    # 使用 PM2 管理进程')

    doc.add_heading('4.4 开发模式启动', level=2)
    doc.add_paragraph('开发模式支持热更新，适合开发调试使用。需要打开 4 个终端分别运行：')
    dev_steps = [
        ('终端 1 — 基础设施', 'docker-compose up -d postgres redis minio'),
        ('终端 2 — 后端（热更新）', 'cd backend && source .venv/bin/activate && uvicorn main:app --host 0.0.0.0 --port 8000 --reload'),
        ('终端 3 — Celery Worker', 'cd backend && source .venv/bin/activate && celery -A app.core.celery_app worker --loglevel=info --concurrency=2'),
        ('终端 4 — 前端（热更新）', 'cd frontend && npm run dev'),
    ]
    add_styled_table(doc, ['终端', '命令'], dev_steps, [2.0, 4.5])

    doc.add_page_break()

    # ========== 5. 配置说明 ==========
    doc.add_heading('5. 配置说明', level=1)

    doc.add_paragraph(
        '系统使用根目录下的 .env 文件作为统一配置源。以下是主要配置项：'
    )

    config_items = [
        ('HOST_IP', 'localhost', '部署主机 IP 或域名'),
        ('BACKEND_PORT', '8000', '后端服务端口'),
        ('FRONTEND_PORT', '3000', '前端服务端口'),
        ('DB_HOST', 'localhost', 'PostgreSQL 主机地址'),
        ('DB_PORT', '5432', 'PostgreSQL 端口'),
        ('DB_USER', 'llmeval', '数据库用户名'),
        ('DB_PASS', 'llmeval123', '数据库密码（生产环境必须修改）'),
        ('DB_NAME', 'llmeval', '数据库名称'),
        ('REDIS_HOST', 'localhost', 'Redis 主机地址'),
        ('REDIS_PORT', '6379', 'Redis 端口'),
        ('MINIO_ENDPOINT', 'localhost:9000', 'MinIO 服务地址'),
        ('MINIO_ACCESS_KEY', 'minioadmin', 'MinIO 访问密钥'),
        ('MINIO_SECRET_KEY', 'minioadmin123', 'MinIO 秘密密钥（生产环境必须修改）'),
        ('SECRET_KEY', 'supersecretkey...', 'JWT 签名密钥（生产环境必须修改）'),
        ('ACCESS_TOKEN_EXPIRE_MINUTES', '1440', 'Token 有效期（分钟，默认 24 小时）'),
        ('CORS_EXTRA_ORIGINS', '', '额外允许的跨域来源'),
        ('HTTP_PROXY / HTTPS_PROXY', '', 'API 代理设置（可选）'),
        ('OPENAI_API_KEY', '', 'OpenAI API 密钥（可在 UI 中逐模型配置）'),
        ('ANTHROPIC_API_KEY', '', 'Anthropic API 密钥（可在 UI 中逐模型配置）'),
    ]
    add_styled_table(doc, ['配置项', '默认值', '说明'], config_items, [2.2, 1.5, 2.8])

    p = doc.add_paragraph()
    run = p.add_run('安全提醒：')
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    p.add_run('生产环境部署时，务必修改 DB_PASS、MINIO_SECRET_KEY、SECRET_KEY 等敏感配置项！')

    doc.add_heading('默认登录凭据', level=2)
    doc.add_paragraph('用户名：admin')
    doc.add_paragraph('密码：admin123')
    p = doc.add_paragraph()
    run = p.add_run('首次登录后请立即修改管理员密码。')
    run.bold = True

    doc.add_page_break()

    # ========== 6. 功能模块详解 ==========
    doc.add_heading('6. 功能模块详解', level=1)

    doc.add_heading('6.1 用户认证与权限', level=2)
    doc.add_paragraph(
        '系统采用 JWT（JSON Web Token）认证机制，支持用户注册、登录和角色管理。'
        '主要包含以下功能：'
    )
    items = [
        '用户注册与登录（JWT Token 认证）',
        '基于角色的权限控制（管理员 / 普通用户）',
        'Token 自动刷新机制',
        '密码加密存储（bcrypt）',
        '团队/组织管理',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.2 模型管理', level=2)
    doc.add_paragraph('在"模型管理"页面（/models）可以添加、编辑和管理待评测的 LLM 模型。')
    doc.add_paragraph('支持的模型提供商：')
    providers = [
        ('OpenAI', 'GPT-4、GPT-4 Turbo、GPT-3.5 Turbo 等'),
        ('Anthropic', 'Claude 3 Opus、Sonnet、Haiku 等'),
        ('DashScope（阿里云）', '通义千问系列'),
        ('DeepSeek', 'DeepSeek 系列模型'),
        ('Zhipu（智谱）', 'GLM 系列模型'),
        ('vLLM', '自托管推理服务'),
        ('Ollama', '本地部署模型'),
        ('自定义 API', '任何兼容 OpenAI 格式的 API'),
    ]
    add_styled_table(doc, ['提供商', '支持模型'], providers, [2.0, 4.5])
    doc.add_paragraph(
        '每个模型可以独立配置 API Key、Base URL、温度参数等。添加模型后可进行连接测试，'
        '确保配置正确。'
    )

    doc.add_heading('6.3 数据集管理', level=2)
    doc.add_paragraph('在"数据集管理"页面（/datasets）可以上传自定义数据集或使用内置数据集。')
    doc.add_paragraph('支持的上传格式：JSONL、JSON、CSV、ZIP')
    doc.add_paragraph('数据集类型：')
    ds_types = [
        '问答对（question + reference_answer）',
        '代码题目（含测试用例）',
        'RAG 数据（含上下文文档）',
        '多轮对话数据',
        '自定义格式',
    ]
    for item in ds_types:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.4 基准测试库', level=2)
    doc.add_paragraph(
        '在"基准测试"页面（/benchmarks）可以浏览和选择内置的标准基准测试。'
        '系统内置 20+ 主流基准测试集：'
    )
    benchmarks = [
        ('MMLU-Pro', '多任务语言理解增强版，涵盖 57 个学科领域'),
        ('GSM8K', '小学数学推理，8000+ 多步骤数学题'),
        ('HumanEval', '代码生成能力评测，164 道编程题'),
        ('C-Eval', '中文知识与能力评测基准'),
        ('TruthfulQA', '事实准确性评测，检测模型幻觉'),
        ('HellaSwag', '常识推理能力评测'),
        ('ARC-Challenge', '科学推理挑战集'),
        ('WinoGrande', '代词消歧义推理'),
        ('MBPP', 'Python 基础编程问题集'),
        ('BBH', 'BIG-Bench Hard 高难推理任务'),
        ('MATH', '高等数学推理'),
        ('DROP', '阅读理解推理'),
        ('HealthBench', '医疗领域专项评测'),
        ('IFEval', '指令遵循评测'),
        ('GPQA Diamond', '研究生级别问答'),
    ]
    add_styled_table(doc, ['基准测试', '描述'], benchmarks, [2.0, 4.5])

    doc.add_heading('6.5 评测任务', level=2)
    doc.add_paragraph(
        '在"评测"页面（/evaluations）可以创建新的评测任务。创建评测的一般流程：'
    )
    steps = [
        '选择一个或多个待评测的模型',
        '选择评测数据集（上传的自定义数据集或内置基准测试）',
        '选择评测维度（如 LLM-as-Judge、基准测试、代码执行、鲁棒性测试等）',
        '设置采样策略和最大样本数',
        '提交评测任务',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_paragraph(
        '评测任务提交后进入异步执行队列（Celery），可在评测详情页（/evaluations/[id]）'
        '通过 WebSocket 实时查看进度。每个模型的评测进度独立追踪，支持取消和重新运行。'
    )

    doc.add_heading('6.6 结果分析与可视化', level=2)
    doc.add_paragraph('评测完成后，在"结果"页面（/results/[id]）查看详细的评测结果：')
    viz_items = [
        '雷达图：多维度能力直观对比',
        '柱状图：单指标横向比较',
        '分布图：得分分布统计',
        '延迟对比：API 响应时间对比',
        '详细指标分解：每个评测维度的细分得分',
        '样本级详情：查看具体的输入/输出和评分明细',
    ]
    for item in viz_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('6.7 报告导出', level=2)
    doc.add_paragraph('系统支持将评测结果导出为多种格式的报告：')
    formats = [
        ('PDF 报告', '包含完整的图表和分析结果，适合分享和存档'),
        ('Excel 报告', '结构化数据表格，便于二次分析'),
        ('JSONL 导出', '原始评测数据，适合程序化处理'),
    ]
    add_styled_table(doc, ['格式', '说明'], formats, [1.5, 5.0])

    doc.add_heading('6.8 ELO 排行榜', level=2)
    doc.add_paragraph(
        'ELO 排行榜（/leaderboard）基于模型对战记录动态计算 ELO 分数。'
        '每次评测完成后自动更新排名，支持按基准测试或评测维度筛选。'
        'ELO 系统参考国际象棋 ELO 等级分制度，能够客观反映模型的相对能力水平。'
    )

    doc.add_heading('6.9 模型竞技场', level=2)
    doc.add_paragraph(
        '模型竞技场（/arena）提供互动式的模型对比功能。用户可以输入问题，'
        '同时获取两个模型的回答，进行并排对比和人工评判。评判结果纳入 ELO 计算。'
    )

    doc.add_heading('6.10 Prompt 管理', level=2)
    doc.add_paragraph(
        'Prompt 管理模块（/prompts）提供 Prompt 模板的创建、版本控制和实验追踪功能。'
        '支持 Prompt 的 A/B 测试，帮助优化评测提示词。'
    )

    doc.add_heading('6.11 AI 助手', level=2)
    doc.add_paragraph(
        '内置 AI 助手可以辅助解读评测结果、提供优化建议。助手基于配置的 Agent 模型运行，'
        '可在设置页面配置 Agent 模型参数。'
    )

    doc.add_heading('6.12 管理后台', level=2)
    doc.add_paragraph(
        '管理后台（/admin）仅管理员可访问，提供以下功能：'
    )
    admin_items = [
        '用户管理：查看、编辑、禁用用户账户',
        '系统审计日志：追踪所有操作记录',
        '平台统计：评测次数、模型数量等全局数据',
    ]
    for item in admin_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ========== 7. 评测维度详解 ==========
    doc.add_heading('7. 评测维度详解', level=1)

    doc.add_paragraph('系统支持以下 20+ 评测维度，可在创建评测任务时灵活组合：')

    dimensions = [
        ('LLM-as-Judge', '使用大语言模型作为评判者，对回答进行多维度打分。支持位置偏差消除、多评判者一致性校验。',
         '评判模型、评分量表、偏差消除策略'),
        ('基准测试', '使用标准基准测试集评估模型能力。支持 MMLU-Pro、GSM8K、HumanEval 等 20+ 基准。',
         '基准名称、采样数量'),
        ('客观指标', '计算 ROUGE、BLEU、METEOR、BERTScore、Exact Match、Token F1、余弦相似度等客观指标。',
         '指标选择'),
        ('幻觉检测', '通过多次采样检测模型回答的一致性，识别潜在幻觉。',
         '采样次数、一致性阈值'),
        ('鲁棒性测试', '对输入进行拼写错误、同义词替换、语序调整等扰动，评估模型稳定性。',
         '扰动类型、扰动强度'),
        ('RAG 评估', '评估检索增强生成的质量，包括上下文相关性、回答忠实度和信息完整性。',
         '上下文数量、评估子维度'),
        ('代码执行', '在 Docker 沙箱中运行模型生成的代码，通过测试用例验证正确性。',
         '编程语言、超时时间、测试用例'),
        ('链式推理 (CoT)', '评估模型 Chain-of-Thought 推理的质量和步骤正确性。',
         '推理步骤评估标准'),
        ('多轮对话', '评估模型在多轮对话中的上下文理解和回答一致性。',
         '对话轮次、上下文长度'),
        ('多语言', '跨语言能力评测，评估模型在不同语言间的翻译和理解能力。',
         '目标语言列表'),
        ('结构化输出', '验证模型输出是否符合指定的 JSON/XML Schema。',
         'Schema 定义、验证规则'),
        ('函数调用', '评估模型的 Function Calling / Tool Use 能力。',
         '函数定义、预期调用'),
        ('长上下文', '测试模型在长文本中检索和引用关键信息的能力。',
         '上下文长度、检索目标'),
        ('指令遵循', '评估模型对格式约束、长度限制等指令的遵守程度。',
         '指令约束列表'),
        ('领域专项', '针对特定领域（医疗、法律、金融等）的专项评测。',
         '领域类型、专业术语'),
        ('性价比分析', 'API 调用延迟、吞吐量分析和成本估算。',
         '计价模型'),
    ]
    for name, desc, params in dimensions:
        doc.add_heading(name, level=3)
        doc.add_paragraph(desc)
        p = doc.add_paragraph()
        run = p.add_run('可配置参数：')
        run.bold = True
        p.add_run(params)

    doc.add_page_break()

    # ========== 8. API 接口概览 ==========
    doc.add_heading('8. API 接口概览', level=1)

    doc.add_paragraph(
        '后端提供完整的 RESTful API，可通过 http://localhost:8000/docs 访问 Swagger 交互文档。'
        '以下是主要 API 路由：'
    )

    apis = [
        ('/api/v1/auth/', '认证', '登录、注册、Token 刷新'),
        ('/api/v1/models/', '模型管理', '模型 CRUD、连接测试'),
        ('/api/v1/datasets/', '数据集', '上传、列表、预览、删除'),
        ('/api/v1/evaluations/', '评测任务', '创建、查询、取消、重新运行'),
        ('/api/v1/benchmarks/', '基准测试', '列表、详情、数据预览'),
        ('/api/v1/results/', '评测结果', '查询结果、获取详情'),
        ('/api/v1/reports/', '报告', '生成、下载（PDF/Excel/JSON）'),
        ('/api/v1/leaderboard/', 'ELO 排行', '排名查询、历史记录'),
        ('/api/v1/arena/', '竞技场', '创建对战、提交评判'),
        ('/api/v1/prompts/', 'Prompt', '模板 CRUD、版本管理'),
        ('/api/v1/metrics/', '指标', '指标定义查询'),
        ('/api/v1/judge-models/', '评判模型', '评判 LLM 配置'),
        ('/api/v1/agent/', 'AI 助手', '对话、流式回复'),
        ('/api/v1/comparison/', '结果对比', '多评测结果对比'),
        ('/api/v1/admin/', '管理', '用户管理、系统统计'),
        ('/api/v1/audit/', '审计', '操作日志查询'),
        ('/api/v1/notifications/', '通知', '通知推送与管理'),
        ('/api/v1/api-keys/', 'API 密钥', '密钥管理'),
        ('/api/v1/teams/', '团队', '团队/组织管理'),
    ]
    add_styled_table(doc, ['路由', '模块', '功能'], apis, [2.0, 1.3, 3.2])

    doc.add_page_break()

    # ========== 9. 常见问题与故障排查 ==========
    doc.add_heading('9. 常见问题与故障排查', level=1)

    faqs = [
        ('无法访问前端页面',
         '检查前端服务是否正常运行（pm2 status 或 docker-compose ps）。'
         '确认 FRONTEND_PORT 未被占用。检查 Nginx 配置是否正确。'),
        ('后端 API 返回 500 错误',
         '查看后端日志（pm2 logs backend 或 docker-compose logs backend）。'
         '检查数据库连接是否正常，确认 .env 中的数据库配置正确。'),
        ('评测任务一直处于等待状态',
         '检查 Celery Worker 是否正常运行。检查 Redis 连接是否正常。'
         '查看 Celery 日志排查错误。'),
        ('模型连接测试失败',
         '确认 API Key 正确且未过期。检查网络连接和代理设置。'
         '如果使用自托管模型，确认推理服务正在运行。'),
        ('MinIO 上传失败',
         '确认 MinIO 服务正常运行。检查 MINIO_ACCESS_KEY 和 MINIO_SECRET_KEY 配置。'
         '确认存储桶已创建（系统启动时自动创建）。'),
        ('数据库迁移错误',
         '运行 alembic upgrade head 手动执行迁移。检查数据库版本兼容性。'
         '确认 PostgreSQL 用户有足够权限。'),
        ('前端页面白屏',
         '清除浏览器缓存，确认 NEXT_PUBLIC_API_URL 配置正确。'
         '检查浏览器控制台是否有 JavaScript 错误。'),
        ('WebSocket 连接失败',
         '确认后端的 WebSocket 端点可达。检查 Nginx 是否正确代理了 WebSocket 连接。'
         '某些防火墙可能阻止 WebSocket 连接。'),
    ]
    for question, answer in faqs:
        doc.add_heading(question, level=3)
        doc.add_paragraph(answer)

    doc.add_page_break()

    # ========== 10. 扩展开发指南 ==========
    doc.add_heading('10. 扩展开发指南', level=1)

    doc.add_heading('添加新的评测器', level=2)
    doc.add_paragraph('按以下步骤添加自定义评测器：')
    steps = [
        '在 eval_engine/evaluators/ 目录下创建新的 Python 文件（如 my_evaluator.py）',
        '实现评测逻辑，返回 0-1 范围的归一化分数',
        '在 eval_engine/engine.py 中注册新评测器',
        '在 backend/app/schemas/evaluation.py 中添加对应的 Schema 字段',
        '在前端评测创建页面中添加新评测维度的 UI 选项',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('添加新的 LLM Provider', level=2)
    doc.add_paragraph('按以下步骤添加新的模型提供商：')
    steps = [
        '在 eval_engine/providers/ 目录下创建新的 Provider 文件',
        '继承 base.py 中的抽象基类，实现 generate() 方法',
        '在 Provider 工厂中注册新的提供商',
        '在前端模型管理页面添加对应的提供商选项',
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f'{i}. {step}')

    doc.add_heading('数据库迁移', level=2)
    doc.add_paragraph(
        '本项目使用 Alembic 管理数据库迁移。修改 ORM 模型后，使用以下命令生成和应用迁移：'
    )
    doc.add_paragraph('alembic revision --autogenerate -m "描述变更"    # 生成迁移脚本')
    doc.add_paragraph('alembic upgrade head    # 应用迁移')

    # ========== 文档尾部 ==========
    doc.add_page_break()
    doc.add_heading('附录', level=1)

    doc.add_heading('访问地址汇总', level=2)
    urls = [
        ('前端界面', 'http://localhost:3000'),
        ('后端 API 文档', 'http://localhost:8000/docs'),
        ('MinIO 控制台', 'http://localhost:9001'),
        ('PostgreSQL', 'localhost:5432'),
        ('Redis', 'localhost:6379'),
    ]
    add_styled_table(doc, ['服务', '地址'], urls, [2.0, 4.5])

    doc.add_heading('技术栈总览', level=2)
    stack = [
        ('前端', 'Next.js 14 + React 18 + TypeScript + Ant Design 5 + ECharts'),
        ('后端', 'FastAPI + Python 3.11 + SQLAlchemy 2 + Celery'),
        ('数据库', 'PostgreSQL 15'),
        ('缓存/队列', 'Redis 7'),
        ('对象存储', 'MinIO'),
        ('反向代理', 'Nginx'),
        ('进程管理', 'PM2 / Docker Compose'),
        ('评测引擎', '20+ 评测器 + 多 LLM Provider 适配'),
    ]
    add_styled_table(doc, ['层级', '技术'], stack, [1.5, 5.0])

    # 保存
    output_path = '/home/liudezhen/Webui-llm-eval/Webui-LLM-Eval_使用说明文档.docx'
    doc.save(output_path)
    print(f'文档已生成：{output_path}')


if __name__ == '__main__':
    main()
